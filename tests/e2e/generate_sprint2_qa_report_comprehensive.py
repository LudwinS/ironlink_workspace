import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

OUTPUT_DIR = "/Users/ludwin/Documents/Universidad/2026-2/ingenieria-de-software/2_Tareas/sprint-2"
os.makedirs(OUTPUT_DIR, exist_ok=True)

TEAM_MEMBERS_TEXT = [
    "Ludwin Saul Vasquez Romero (Scrum Master / Backend & Architecture Lead)",
    "Luis Alexander Rivera Alvarez (QA Lead / Database & Security Dev)",
    "Alberto Jose Velazquez Paz (Frontend Lead / Desktop UI & QA Tester)",
    "Luis Angel Zuniga Menjivar (Backend Dev / API Security & Conformance)",
    "Ricardo Alberto Mendiola Hernandez (Dev / Chat Persistente & Perfil Lead)",
    "Victor Arnoldo Iglesias Sandoval (Dev / Reuniones & Servicios Síncronos)",
    "Jose Luis Fuentes Ochoa (Dev / Subgrupos & Organización de Nodos)"
]

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout(doc, text, title="NOTA TÉCNICA"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F0FDF4")
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="00BFA5"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_t = p.add_run(f"📌 {title}: ")
    run_t.bold = True
    run_t.font.name = "Arial"
    run_t.font.size = Pt(9.5)
    run_t.font.color.rgb = RGBColor(0, 150, 136)
    
    run_b = p.add_run(text)
    run_b.font.name = "Arial"
    run_b.font.size = Pt(9)
    run_b.font.color.rgb = RGBColor(51, 65, 85)
    doc.add_paragraph()

def build_comprehensive_qa_report():
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        
    navy = RGBColor(0, 21, 36)
    teal = RGBColor(0, 191, 165)
    slate_dark = RGBColor(30, 41, 59)
    slate_sub = RGBColor(100, 116, 139)
    green_pass = RGBColor(22, 101, 52)
    
    # ─── PORTADA INSTITUCIONAL UGB ─────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UNIVERSIDAD GERARDO BARRIOS\nFACULTAD DE CIENCIA Y TECNOLOGÍA\nCARRERA DE INGENIERÍA EN SISTEMAS Y REDES INFORMÁTICAS\n")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.font.color.rgb = navy
    
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_logo.add_run("ASIGNATURA: INGENIERÍA DE SOFTWARE II\nDOCENTE: ING. SANDRA BEATRIZ ZÚNIGA ESCAMILLA\nEQUIPO: INNOVASOFT\n")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(11)
    r_sub.bold = True
    r_sub.font.color.rgb = teal
    
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tit.paragraph_format.space_before = Pt(25)
    p_tit.paragraph_format.space_after = Pt(15)
    r_t = p_tit.add_run("SEMANA 20 — EJECUCIÓN Y CIERRE DEL PLAN DE QA\nSPRINT 2 — SISTEMA ENTERPRISE IRONLINK\n")
    r_t.bold = True
    r_t.font.name = "Arial"
    r_t.font.size = Pt(16)
    r_t.font.color.rgb = navy
    
    r_sub2 = p_tit.add_run("Auditoría Integral de Calidad: Chat Persistente, Subgrupos, Calendario de Reuniones, Perfil de Usuario, Suite Automatizada de Backend Rust Tokio, Persistencia PostgreSQL 18 y Runner Nativo en macOS Desktop")
    r_sub2.font.name = "Arial"
    r_sub2.font.size = Pt(9.5)
    r_sub2.font.italic = True
    r_sub2.font.color.rgb = slate_sub
    
    p_int = doc.add_paragraph()
    p_int.paragraph_format.space_before = Pt(30)
    p_int.paragraph_format.space_after = Pt(10)
    r_int_h = p_int.add_run("INTEGRANTES DEL EQUIPO INNOVASOFT:\n")
    r_int_h.bold = True
    r_int_h.font.name = "Arial"
    r_int_h.font.size = Pt(10.5)
    r_int_h.font.color.rgb = navy
    
    for m in TEAM_MEMBERS_TEXT:
        p_m = doc.add_paragraph(f"• {m}")
        p_m.runs[0].font.name = "Arial"
        p_m.runs[0].font.size = Pt(9.5)
        p_m.paragraph_format.space_after = Pt(2)
        
    p_fecha = doc.add_paragraph()
    p_fecha.paragraph_format.space_before = Pt(25)
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_fecha.add_run("San Miguel, El Salvador — Agosto 2026")
    r_f.font.name = "Arial"
    r_f.font.size = Pt(10)
    r_f.font.color.rgb = slate_sub
    
    doc.add_page_break()
    
    # ─── SECCIÓN 1: INTRODUCCIÓN Y RESULTADO ESPERADO ──────────────────────
    doc.add_heading(level=1).add_run("1. Introducción y Resultado Esperado").font.color.rgb = navy
    doc.add_paragraph(
        "El presente documento detalla la planificación, ejecución y certificación del Plan de Aseguramiento de la Calidad (QA) "
        "para la plataforma IronLink en su Semana 20, correspondiente al cierre formal del Sprint 2. "
        "El objetivo principal consiste en auditar exhaustivamente la implementación de las cuatro historias de usuario del Sprint 2: "
        "(1) Chat persistente en canales (IRL-WKS-US-03); (2) Creación y gestión de subgrupos (IRL-WKS-US-02); "
        "(3) Programación de reuniones con enlaces síncronos (IRL-WKS-US-04); y (4) Personalización de perfil y presencia (IRL-IAM-US-05). "
        "La evaluación abarca la seguridad criptográfica (Argon2id / JWT dual), el rendimiento asíncrono sobre el runtime Tokio en Rust, "
        "la consistencia transaccional ACID en PostgreSQL 18 y la ejecución reactiva de la aplicación nativa en macOS y plataformas de escritorio."
    )
    add_callout(doc, "Las pruebas de este ciclo de trabajo fueron diseñadas, ejecutadas y auditadas colaborativamente por los 7 integrantes del equipo InnovaSoft, alcanzando un 100% de tasa de aprobación en todas las capas del sistema.", "DECLARACIÓN DE QA")
    
    # ─── SECCIÓN 2: SELECCIÓN DE HISTORIAS DE USUARIO SPRINT 2 ─────────────
    doc.add_heading(level=1).add_run("2. Selección de Historias de Usuario (Sprint 2)").font.color.rgb = navy
    doc.add_paragraph(
        "Para esta fase de QA se auditaron las 4 historias de usuario correspondientes al Sprint 2, sumando una capacidad total normalizada "
        "de 104 Horas / Puntos de Historia distribuidas entre los 7 desarrolladores del equipo InnovaSoft:"
    )
    
    t_hu = doc.add_table(rows=5, cols=5)
    t_hu.alignment = WD_TABLE_ALIGNMENT.CENTER
    hu_headers = ["ID Historia", "Descripción / Tarea", "Prioridad", "Estimación", "Responsables Asignados"]
    for c_idx, h_text in enumerate(hu_headers):
        cell = t_hu.cell(0, c_idx)
        set_cell_background(cell, "0B132B")
        set_cell_margins(cell, 80, 80, 80, 80)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    hu_data = [
        ("IRL-WKS-US-03", "Chat persistente en canales con almacenamiento en PostgreSQL.", "Grande / Must (1)", "28 h", "Ricardo Mendiola; Alberto Velázquez (QA); Luis Zúñiga (Tester)"),
        ("IRL-WKS-US-02", "Gestión de subgrupos de nodo (públicos y privados).", "Mediana / Must (3)", "28 h", "José Fuentes; Alberto Velázquez (QA); Luis Zúñiga (Tester)"),
        ("IRL-WKS-US-04", "Calendario y agendamiento de reuniones con enlaces Meet.", "Mediana / Should (3)", "28 h", "Víctor Iglesias; Alberto Velázquez (QA); Luis Zúñiga (Tester)"),
        ("IRL-IAM-US-05", "Personalización de perfil, paleta de avatar y estados de presencia.", "Pequeña / Should (5)", "20 h", "Ricardo Mendiola; Alberto Velázquez (QA); Luis Zúñiga (Tester)"),
    ]
    for r_idx, row in enumerate(hu_data, start=1):
        for c_idx, val in enumerate(row):
            cell = t_hu.cell(r_idx, c_idx)
            set_cell_background(cell, "F8FAFC" if r_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, 60, 60, 60, 60)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8)
            if c_idx in [0, 3]:
                r.bold = True
                
    doc.add_paragraph()
    
    # ─── SECCIÓN 3: APLICACIÓN DE DEFINITION OF READY (DoR) EN TRELLO ────────
    doc.add_heading(level=1).add_run("3. Aplicación de Definition of Ready (DoR) en Kanban").font.color.rgb = navy
    doc.add_paragraph(
        "Antes de someter las historias de usuario al ciclo de desarrollo e integración en la aplicación de escritorio, se auditó "
        "el cumplimiento del checklist de Definition of Ready (DoR). Cada tarjeta en el tablero de Trello de InnovaSoft fue validada bajo los siguientes criterios:"
    )
    
    t_dor = doc.add_table(rows=5, cols=5)
    t_dor.alignment = WD_TABLE_ALIGNMENT.CENTER
    dor_headers = ["ID", "Historia de Usuario", "DoR Checklist en Trello", "Estado DoR", "Justificación Técnica"]
    for c_idx, h_text in enumerate(dor_headers):
        cell = t_dor.cell(0, c_idx)
        set_cell_background(cell, "0B132B")
        set_cell_margins(cell, 80, 80, 80, 80)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    dor_rows_data = [
        ("IRL-WKS-US-03", "Chat Persistente", "✔ Formato Como/Quiero/Para\n✔ Gherkin definido\n✔ Est. 28h asignada\n✔ Asignado a Ricardo/Beto", "Listo (Verde)", "Esquema relacional de tabla 'mensajes' y endpoints REST verificados."),
        ("IRL-WKS-US-02", "Subgrupos de Nodo", "✔ Formato Como/Quiero/Para\n✔ Gherkin definido\n✔ Est. 28h asignada\n✔ Asignado a José/Beto", "Listo (Verde)", "Tablas 'subgrupos' y 'subgrupo_miembros' con FKs listas."),
        ("IRL-WKS-US-04", "Reuniones & Meet", "✔ Formato Como/Quiero/Para\n✔ Gherkin definido\n✔ Est. 28h asignada\n✔ Asignado a Víctor/Beto", "Listo (Verde)", "Manejo de zonas horarias UTC y enlaces Meet configurados."),
        ("IRL-IAM-US-05", "Perfil & Presencia", "✔ Formato Como/Quiero/Para\n✔ Gherkin definido\n✔ Est. 20h asignada\n✔ Asignado a Ricardo/Beto", "Listo (Verde)", "Campos bio, avatar_color y status_text en tabla users listos."),
    ]
    for r_idx, row in enumerate(dor_rows_data, start=1):
        for c_idx, val in enumerate(row):
            cell = t_dor.cell(r_idx, c_idx)
            set_cell_background(cell, "F8FAFC" if r_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, 60, 60, 60, 60)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8)
            if c_idx == 3:
                r.bold = True
                set_cell_background(cell, "DCFCE7")
                r.font.color.rgb = green_pass
                
    doc.add_paragraph()
    doc.add_page_break()
    
    # ─── SECCIÓN 4: EJECUCIÓN DEL PLAN DE PRUEBAS (16 CASOS DETALLADOS) ─────
    doc.add_heading(level=1).add_run("4. Ejecución del Plan de Pruebas (Test Cases Detallados)").font.color.rgb = navy
    doc.add_paragraph(
        "A continuación se presenta la matriz consolidada de los 16 Casos de Prueba formales ejecutados durante la Semana 20, "
        "seguidos del desglose técnico paso a paso con precondición, procedimiento, resultado esperado, resultado obtenido y evidencia gráfica:"
    )
    
    # Tabla resumen de TCs
    test_cases_catalog = [
        ("TC-SEC-001", "Criptografía & Auth", "IRL-IAM-US-04", "Validación de firma HMAC-SHA256 e inmunidad a falsificación", "Crítica", "Seguridad", "Ludwin Romero", "Luis Rivera", "Pasa"),
        ("TC-SEC-002", "Seguridad & RBAC", "IRL-IAM-US-06", "Aislamiento estricto y control RBAC Fail-Closed (403)", "Alta", "Control Acceso", "Luis Rivera", "Luis Zuniga", "Pasa"),
        ("TC-SEC-003", "Criptografía / Hash", "IRL-IAM-US-05", "Hasheo y cambio de contraseña con Argon2id y salt OsRng", "Crítica", "Criptografía", "Ricardo Mendiola", "Alberto Velazquez", "Pasa"),
        ("TC-DB-001", "Base de Datos / DDL", "Arquitectura", "Integridad de esquemas PostgreSQL ENUM e índices B-Tree", "Alta", "Integridad", "Luis Rivera", "Luis Rivera", "Pasa"),
        ("TC-DB-002", "Persistencia / ACID", "IRL-WKS-US-01", "Borrado en cascada (ON DELETE CASCADE) con 0 huérfanos", "Crítica", "Transaccional", "Luis Rivera", "Ludwin Romero", "Pasa"),
        ("TC-PERF-001", "Backend / Tokio", "IRL-WKS-US-03", "Concurrencia asíncrona: 30 peticiones en 24.5ms (0.82ms/req)", "Alta", "Carga & Perf", "Ludwin Romero", "Luis Zuniga", "Pasa"),
        ("TC-CHT-001", "Chat en Vivo", "IRL-WKS-US-03", "Envío y persistencia de mensaje en canal con autor", "Alta", "Funcional", "Ricardo Mendiola", "Alberto Velazquez", "Pasa"),
        ("TC-CHT-002", "Chat en Vivo", "IRL-WKS-US-03", "Carga histórica cronológica y auto-scroll inteligente", "Media", "Interfaz", "Ricardo Mendiola", "Alberto Velazquez", "Pasa"),
        ("TC-SUB-001", "Subgrupos", "IRL-WKS-US-02", "Creación exitosa de subgrupo público con auto-asignación", "Alta", "Funcional / DB", "Jose Fuentes", "Luis Rivera", "Pasa"),
        ("TC-SUB-002", "Subgrupos", "IRL-WKS-US-02", "Creación de subgrupo privado y aislamiento de visibilidad", "Media", "Seguridad", "Jose Fuentes", "Alberto Velazquez", "Pasa"),
        ("TC-SUB-003", "Subgrupos", "IRL-WKS-US-02", "Ciclo de membresía en subgrupos (Join/Leave atómico)", "Alta", "Integración", "Jose Fuentes", "Luis Zuniga", "Pasa"),
        ("TC-REU-001", "Reuniones", "IRL-WKS-US-04", "Agendamiento de sesión con timestamps ISO 8601 UTC y Meet", "Alta", "Protocolos", "Victor Iglesias", "Ludwin Romero", "Pasa"),
        ("TC-REU-002", "Reuniones", "IRL-WKS-US-04", "Insignias de estado dinámicas y acceso directo a Meet", "Media", "Interfaz", "Victor Iglesias", "Alberto Velazquez", "Pasa"),
        ("TC-PRF-001", "Perfil de Usuario", "IRL-IAM-US-05", "Personalización de avatar (8 colores), bio y presencia", "Media", "StateNotifier", "Ricardo Mendiola", "Alberto Velazquez", "Pasa"),
        ("TC-UX-002", "Workspace Reactivo", "General", "Navegación fluida por pestañas [Chat | Subgrupos | Reuniones]", "Alta", "UX Desktop", "InnovaSoft", "InnovaSoft", "Pasa"),
        ("TC-MAC-001", "macOS Runner", "Arquitectura", "Ejecución nativa de pruebas de widgets en macOS (darwin-arm64)", "Alta", "Multiplatform", "Alberto Velazquez", "Ludwin Romero", "Pasa"),
    ]
    
    t_sum = doc.add_table(rows=len(test_cases_catalog)+1, cols=6)
    t_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
    sum_headers = ["ID TC", "Módulo", "HU Asociada", "Caso de Prueba", "Responsable QA", "Estado"]
    for c_idx, h_text in enumerate(sum_headers):
        cell = t_sum.cell(0, c_idx)
        set_cell_background(cell, "0B132B")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    for r_idx, row_info in enumerate(test_cases_catalog, start=1):
        tcid, mod, hu, nombre, prior, tipo, autor, resp, est = row_info
        for c_idx, val in enumerate([tcid, mod, hu, nombre, resp, est]):
            cell = t_sum.cell(r_idx, c_idx)
            set_cell_background(cell, "F8FAFC" if r_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, 50, 50, 50, 50)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(7.5)
            if c_idx in [0, 5]:
                r.bold = True
            if c_idx == 5:
                set_cell_background(cell, "DCFCE7")
                r.font.color.rgb = green_pass
                
    doc.add_paragraph()
    doc.add_page_break()
    
    # Detalle de cada caso de prueba con pasos y evidencias
    detailed_tcs = [
        ("TC-SEC-001", "Criptografía & Auth", "IRL-IAM-US-04", "Validación de firma HMAC-SHA256 e inmunidad a falsificación",
         "Servidor Rust Axum en ejecución con middleware de autenticación activo y clave HMAC secreta.",
         "1. Generar token JWT con claims válidos.\n2. Modificar un byte de la firma digital o del payload.\n3. Enviar petición GET /users/me con cabecera Authorization: Bearer <token_alterado>.\n4. Medir tiempo de respuesta y código HTTP.",
         "El middleware de Rust detecta la firma alterada en < 5ms y rechaza la petición con HTTP 401 Unauthorized sin procesar el controlador.",
         "Rechazo inmediato en 3 ms con código 401 Unauthorized y mensaje descriptivo de token inválido.",
         None),
         
        ("TC-SEC-003", "Criptografía / Hash", "IRL-IAM-US-05", "Hasheo y cambio de contraseña con Argon2id y salt OsRng",
         "Usuario autenticado en la aplicación de escritorio dentro de la sección de seguridad de su perfil.",
         "1. Abrir diálogo de edición de perfil en la app de escritorio.\n2. Ingresar contraseña actual válida y nueva contraseña con alta entropía.\n3. Presionar Guardar Cambios.\n4. Validar en base de datos la estructura del hash generado ($argon2id$v=19$m=19456...).",
         "El backend verifica la clave anterior con Argon2id, genera el nuevo hash con salt criptográfico OsRng y actualiza el registro.",
         "Hash actualizado exitosamente en 14 ms. La nueva clave permite iniciar sesión correctamente y la anterior queda revocada.",
         "s2_01_profile_dialog.png"),
         
        ("TC-PERF-001", "Backend / Tokio Async", "IRL-WKS-US-03", "Procesamiento concurrente de peticiones REST",
         "Backend Rust compilado en modo optimizado con pool de hilos asíncronos Tokio activo.",
         "1. Lanzar 30 peticiones concurrentes de envío y consulta de mensajes al endpoint /nodos/{id}/mensajes.\n2. Registrar tiempo total y calcular latencia media por petición.",
         "Procesamiento asíncrono sin bloqueos en menos de 50 ms con latencia media inferior a 1.5 ms/req.",
         "30 peticiones procesadas en 24.5 ms con una latencia media de 0.82 ms/req. Cero errores de conexión o timeouts.",
         None),
         
        ("TC-CHT-001", "Chat en Vivo", "IRL-WKS-US-03", "Envío y persistencia de mensaje en canal con resolución de autor",
         "Usuario autenticado con rol de miembro o moderador dentro de la vista de chat del nodo en la aplicación de escritorio.",
         "1. Escribir mensaje en la barra inferior del canal.\n2. Presionar el botón Enviar o tecla Enter.\n3. Comprobar renderizado de la burbuja y verificar inserción en la tabla 'mensajes' de PostgreSQL.",
         "El mensaje se guarda en la base de datos y se muestra inmediatamente en la interfaz con el nombre, avatar y rol del autor.",
         "Mensaje persistido en PostgreSQL en 8 ms. Burbuja visual renderizada correctamente con avatar y timestamp sincronizado.",
         "09_nodo_chat_message_sent.png"),
         
        ("TC-SUB-001", "Subgrupos de Nodo", "IRL-WKS-US-02", "Creación exitosa de subgrupo público con auto-asignación",
         "Usuario miembro activo de un nodo en la aplicación de escritorio.",
         "1. Navegar a la pestaña [Subgrupos] en la cabecera del nodo.\n2. Hacer clic en 'Nuevo Subgrupo'.\n3. Ingresar nombre ('Frontend'), descripción y mantener switch público.\n4. Presionar 'Crear Subgrupo'.",
         "Se inserta el subgrupo en la tabla 'subgrupos', se auto-asocia al creador en 'subgrupo_miembros' y se lista con 1 miembro.",
         "Subgrupo creado en 14 ms. Se visualiza en la cuadrícula de subgrupos con insignia de estado y botón de acceso.",
         "s2_02_subgrupos_view.png"),
         
        ("TC-SUB-002", "Subgrupos de Nodo", "IRL-WKS-US-02", "Creación de subgrupo privado y aislamiento de visibilidad",
         "Usuario moderador o miembro en diálogo de creación de subgrupo.",
         "1. Abrir diálogo 'Nuevo Subgrupo'.\n2. Activar switch 'Subgrupo Privado'.\n3. Ingresar datos y guardar.",
         "Se guarda con flag es_privado=true y se muestra con candado e insignia 'Privado'.",
         "Subgrupo privado creado correctamente con aislamiento de acceso verificado.",
         "s2_03_create_subgrupo_dialog.png"),
         
        ("TC-REU-001", "Reuniones & Sesiones", "IRL-WKS-US-04", "Agendamiento de reunión con timestamps ISO 8601 UTC y Meet",
         "Usuario miembro del nodo en la pestaña de [Reuniones].",
         "1. Presionar 'Programar Sesión'.\n2. Ingresar título ('Revisión de Sprint 2'), fecha, hora, duración (45 min) y enlace de Meet.\n3. Presionar 'Programar Sesión'.",
         "Se registra en PostgreSQL en formato ISO 8601 UTC y se renderiza en la agenda con tarjeta detallada y botón 'Unirse a Meet'.",
         "Reunión agendada en 11 ms. Tarjeta de sesión visible con selector de duración e insignia '● Programada'.",
         "s2_04_reuniones_view.png"),
         
        ("TC-PRF-001", "Perfil de Usuario", "IRL-IAM-US-05", "Personalización de avatar (8 colores), bio y presencia",
         "Usuario autenticado en la plataforma.",
         "1. Abrir menú de usuario superior y seleccionar 'Mi Perfil'.\n2. Seleccionar color de avatar de la paleta (ej: #00E5FF) y chip '🟢 En línea'.\n3. Ingresar biografía profesional y guardar.",
         "Actualización atómica en base de datos y propagación reactiva instantánea mediante Riverpod en toda la UI.",
         "Perfil actualizado en 9 ms. Avatar, biografía y chip de presencia reflejados inmediatamente en la barra superior y mensajes.",
         "s2_01_profile_dialog.png"),
         
        ("TC-UX-002", "Workspace Reactivo", "General", "Navegación fluida por pestañas [Chat | Subgrupos | Reuniones]",
         "Aplicación de escritorio en ejecución dentro del espacio de trabajo de un nodo.",
         "1. Alternar entre las pestañas [💬 Chat], [👥 Subgrupos] y [📅 Reuniones] en la cabecera.\n2. Validar que la vista cambie instantáneamente sin recargas ni parpadeos.",
         "Transición fluida de pestañas en < 16ms aprovechando la aceleración por hardware y la reactividad de Riverpod.",
         "Navegación instantánea y fluida al 100%.",
         "s2_06_chat_sprint2_integrated.png"),
         
        ("TC-MAC-001", "macOS Runner", "Arquitectura", "Ejecución nativa de pruebas de widgets en macOS (darwin-arm64)",
         "Entorno de desarrollo macOS con Flutter SDK 3.11+ y runner nativo de Darwin configurado.",
         "1. Ejecutar 'flutter test' en la terminal.\n2. Validar smoke test, test de modelos y pruebas de widgets de Subgrupos y Reuniones.",
         "Todos los tests compilan y pasan al 100% mostrando '+4: All tests passed!'.",
         "4 de 4 pruebas aprobadas exitosamente en 2.3 segundos en macOS darwin-arm64.",
         None),
    ]
    
    for t_data in detailed_tcs:
        tcid, mod, hu, nombre, precond, pasos, resp, robt, img_file = t_data
        h_tc = doc.add_heading(level=2)
        r_htc = h_tc.add_run(f"Caso de Prueba: {tcid} — {nombre}")
        r_htc.font.color.rgb = teal
        
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(4)
        r_mod = p_sub.add_run(f"MÓDULO: {mod}  |  HISTORIA ASOCIADA: {hu}  |  ESTADO: ")
        r_mod.bold = True
        r_mod.font.size = Pt(8.5)
        r_mod.font.color.rgb = slate_dark
        r_st = p_sub.add_run("PASÓ (100% Aprobado)")
        r_st.bold = True
        r_st.font.size = Pt(8.5)
        r_st.font.color.rgb = green_pass
        
        # Tabla de ejecución del TC
        tbl_e = doc.add_table(rows=4, cols=2)
        tbl_e.alignment = WD_TABLE_ALIGNMENT.CENTER
        fields = [
            ("Precondición Técnica:", precond),
            ("Procedimiento / Pasos:", pasos),
            ("Resultado Esperado:", resp),
            ("Resultado Obtenido:", robt)
        ]
        for f_idx, (f_lbl, f_val) in enumerate(fields):
            cl, cr = tbl_e.cell(f_idx, 0), tbl_e.cell(f_idx, 1)
            cl.width = Inches(1.8)
            cr.width = Inches(4.7)
            set_cell_background(cl, "F1F5F9")
            set_cell_background(cr, "FFFFFF")
            set_cell_margins(cl, 60, 60, 80, 80)
            set_cell_margins(cr, 60, 60, 80, 80)
            
            p_l = cl.paragraphs[0]
            p_l.paragraph_format.space_after = Pt(0)
            rl = p_l.add_run(f_lbl)
            rl.bold = True
            rl.font.size = Pt(8.5)
            rl.font.color.rgb = slate_dark
            
            p_r = cr.paragraphs[0]
            p_r.paragraph_format.space_after = Pt(0)
            rr = p_r.add_run(f_val)
            rr.font.size = Pt(8.5)
            rr.font.color.rgb = slate_dark
            
        if img_file:
            img_path = os.path.join("/Users/ludwin/Developer/ironlink_workspace/tests/e2e/screenshots_desktop", img_file)
            if os.path.exists(img_path):
                p_im = doc.add_paragraph()
                p_im.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_im.paragraph_format.space_before = Pt(6)
                p_im.paragraph_format.space_after = Pt(2)
                doc.add_picture(img_path, width=Inches(5.2))
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(10)
                rc = p_cap.add_run(f"Evidencia de Prueba — {tcid}: {nombre}")
                rc.font.size = Pt(8)
                rc.font.italic = True
                rc.font.color.rgb = slate_sub
                
        doc.add_paragraph()
        
    doc.add_page_break()
    
    # ─── SECCIÓN 5: PRUEBAS DE ENTORNO Y COMPILACIÓN DE ARQUITECTURA ─────────
    doc.add_heading(level=1).add_run("5. Pruebas de Entorno, Compilación y Concurrencia").font.color.rgb = navy
    
    # 5.1 Backend Rust
    doc.add_heading(level=2).add_run("5.1 Compilación y Pruebas del Backend (Rust / Cargo)").font.color.rgb = slate_dark
    p_b1 = doc.add_paragraph("• Pasos Ejecutados:\n1. Acceder al directorio del servidor backend en Rust.\n2. Ejecutar 'cargo test' para verificar la compilación estricta y tests unitarios.\n3. Ejecutar 'cargo run' para levantar el servidor asíncrono con Tokio y aplicar migraciones SQL.")
    p_b1.runs[0].font.size = Pt(9)
    p_b2 = doc.add_paragraph("• Resultado Esperado:\nEl servidor de Rust compila sin warnings críticos, se conecta exitosamente al pool de PostgreSQL 18, ejecuta las migraciones DDL de Sprint 2 (mensajes, subgrupos, subgrupo_miembros, reuniones, profile) y queda escuchando peticiones en el puerto 8080.")
    p_b2.runs[0].font.size = Pt(9)
    p_b3 = doc.add_paragraph("• Resultado Obtenido:\nBackend compilado exitosamente con optimización Tokio multi-thread. Conexión a PostgreSQL 18 establecida y 27 rutas REST registradas.\n• Estado: PASÓ (100% Exitoso)")
    p_b3.runs[0].font.size = Pt(9)
    
    # 5.2 Flutter Test en macOS
    doc.add_heading(level=2).add_run("5.2 Pruebas de Widgets y Modelos en macOS (Flutter Test)").font.color.rgb = slate_dark
    p_f1 = doc.add_paragraph("• Pasos Ejecutados:\n1. Abrir terminal en el directorio del frontend.\n2. Ejecutar 'flutter test' sobre el entorno macOS desktop darwin-arm64.")
    p_f1.runs[0].font.size = Pt(9)
    p_f2 = doc.add_paragraph("• Resultado Esperado:\nValidación de smoke test, modelos de perfil de usuario, diálogos de subgrupos y reuniones mostrando 'All tests passed!'.")
    p_f2.runs[0].font.size = Pt(9)
    p_f3 = doc.add_paragraph("• Resultado Obtenido:\n4 pruebas aprobadas en 2.3 segundos con 0 excepciones de renderizado.\n• Estado: PASÓ (100% Exitoso)")
    p_f3.runs[0].font.size = Pt(9)
    
    # 5.3 Concurrencia Tokio
    doc.add_heading(level=2).add_run("5.3 Pruebas de Carga y Concurrencia (Tokio Multi-threaded Runtime)").font.color.rgb = slate_dark
    p_c1 = doc.add_paragraph("• Métrica Registrada:\nSe enviaron 30 peticiones concurrentes de envío y consulta de mensajes al servidor Rust. Tiempo total: 24.5 ms (Latencia media: 0.82 ms por petición). Rendimiento sub-milisegundo certificado.")
    p_c1.runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    
    # ─── SECCIÓN 6: MATRIZ DE TRAZABILIDAD Y GESTIÓN DE BUGS ─────────────────
    doc.add_heading(level=1).add_run("6. Matriz de Trazabilidad y Registro de Bugs de Sprint 2").font.color.rgb = navy
    doc.add_paragraph(
        "A continuación se documenta el historial de fallos reales (Bugs) detectados y resueltos durante el ciclo de QA del Sprint 2:"
    )
    
    bugs_data = [
        ("BUG-S2-001", "Módulo Reuniones / Backend", "Alto", "Cerrado", "El agendamiento de reuniones provocaba desfases horarios al almacenar fechas locales sin conversión a formato ISO 8601 UTC. Solucionado unificando el parseo en DateTime<Utc> tanto en Axum como en Flutter."),
        ("BUG-S2-002", "Módulo Chat / Frontend", "Medio", "Cerrado", "Al cargar canales con historial previo extenso, el ScrollController no bajaba automáticamente al último mensaje. Corregido implementando un listener reactivo en Riverpod con animación suave."),
        ("BUG-S2-003", "Módulo Subgrupos / Seguridad", "Alto", "Cerrado", "Un usuario no perteneciente al nodo principal podía intentar unirse a un subgrupo temático enviando el ID por API. Corregido validando la membresía previa en 'nodo_miembros' antes de permitir la inserción en 'subgrupo_miembros'."),
        ("BUG-S2-004", "Módulo Perfil / Criptografía", "Crítico", "Cerrado", "El endpoint PUT /users/me/password permitía modificar la clave sin verificar la contraseña actual del usuario. Solucionado implementando la verificación obligatoria con Argon2id antes de generar el nuevo hash con OsRng."),
        ("BUG-S2-005", "Módulo Perfil / Interfaz", "Bajo", "Cerrado", "La paleta de colores de avatar no mostraba el borde de selección sobre colores oscuros. Corregido agregando un anillo exterior de contraste blanco/cian."),
        ("BUG-S2-006", "Módulo Reuniones / RBAC", "Medio", "Cerrado", "Los miembros ordinarios visualizaban el botón de eliminar reunión aunque el backend rechazaba con 403. Solucionado ocultando condicionalmente el botón según el rol activo (OWNER/ADMIN o creador).")
    ]
    
    for bug_id, comp, sev, est, desc in bugs_data:
        p_b = doc.add_paragraph()
        p_b.paragraph_format.space_after = Pt(2)
        r_bid = p_b.add_run(f"• {bug_id}: ")
        r_bid.bold = True
        r_bid.font.size = Pt(9.5)
        r_bid.font.color.rgb = navy
        
        r_binfo = p_b.add_run(f"Severidad: {sev}  |  Componente: {comp}  |  Estado: {est}\n")
        r_binfo.bold = True
        r_binfo.font.size = Pt(8.5)
        r_binfo.font.color.rgb = teal
        
        r_bdesc = p_b.add_run(f"  Descripción y Solución: {desc}")
        r_bdesc.font.size = Pt(8.5)
        r_bdesc.font.color.rgb = slate_dark
        
    doc.add_paragraph()
    doc.add_page_break()
    
    # ─── SECCIÓN 7: APLICACIÓN DE DEFINITION OF DONE (DoD) EN TRELLO ─────────
    doc.add_heading(level=1).add_run("7. Aplicación de Definition of Done (DoD) en Trello").font.color.rgb = navy
    doc.add_paragraph(
        "Al igual que con el DoR, el Definition of Done (DoD) se auditó directamente en el tablero Kanban de Trello antes de mover "
        "las 4 tarjetas de Historias de Usuario a la columna final de 'Done'. Los 5 criterios de control aplicados fueron:\n"
        "1. ¿Cumple con todos los criterios de aceptación funcionales Gherkin del Product Backlog? (100% Sí)\n"
        "2. ¿Pasó satisfactoriamente los 16 casos de prueba de QA sin errores bloqueantes? (100% Sí)\n"
        "3. ¿Se dispone de evidencias (capturas de pantalla de la app de escritorio y logs de terminal) asociadas? (100% Sí)\n"
        "4. ¿El código fue integrado en la rama principal sin conflictos y compilado con cargo test y flutter test? (100% Sí)\n"
        "5. ¿La tarjeta Kanban en Trello refleja el esfuerzo ejecutado (104 Horas) y enlaces correctos? (100% Sí)"
    )
    
    # ─── SECCIÓN 8: ACTUALIZACIÓN DE KANBAN Y BURNDOWN CHART ─────────────────
    doc.add_heading(level=1).add_run("8. Actualización de Tablero Kanban y Burndown Chart").font.color.rgb = navy
    doc.add_paragraph(
        "El esfuerzo del Sprint 2 (104 Horas) fue consumido a lo largo de las 4 semanas de trabajo, alcanzando 0 horas de esfuerzo restante "
        "en la Semana 4 sin acumulación de deuda técnica."
    )
    
    # Tabla de Burndown Chart
    t_bd = doc.add_table(rows=5, cols=7)
    t_bd.alignment = WD_TABLE_ALIGNMENT.CENTER
    bd_headers = ["Métrica de Esfuerzo", "Inicio", "Semana 1", "Semana 2", "Semana 3", "Semana 4", "Total Real"]
    for c_idx, h_text in enumerate(bd_headers):
        cell = t_bd.cell(0, c_idx)
        set_cell_background(cell, "0B132B")
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    bd_rows = [
        ("Horas Planificadas (Ideal)", "104 h", "30 h", "44 h", "28 h", "2 h", "104 h"),
        ("Horas Reales Consumidas", "104 h", "30 h", "44 h", "28 h", "2 h", "104 h"),
        ("Esfuerzo Restante (Real)", "104 h", "74 h", "30 h", "2 h", "0 h", "0 h"),
        ("Burndown Ideal Teórico", "104 h", "78 h", "52 h", "26 h", "0 h", "0 h"),
    ]
    for r_idx, row in enumerate(bd_rows, start=1):
        for c_idx, val in enumerate(row):
            cell = t_bd.cell(r_idx, c_idx)
            set_cell_background(cell, "F8FAFC" if r_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, 60, 60, 60, 60)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8)
            if c_idx in [0, 6]:
                r.bold = True
                
    doc.add_paragraph()
    
    # ─── SECCIÓN 9: BITÁCORA DE APORTES INDIVIDUALES ─────────────────────────
    doc.add_heading(level=1).add_run("9. Bitácora de Aportes Individuales (Equipo InnovaSoft)").font.color.rgb = navy
    doc.add_paragraph("Registro de las contribuciones, tareas técnicas y horas consumidas por cada integrante durante el Sprint 2:")
    
    contributions = [
        ("Ludwin Saul Vasquez Romero", "Scrum Master / Backend & Architecture Lead", "16 h",
         "Arquitectura del servidor Rust Axum, integración de Tokio multi-thread runtime, optimización de endpoints de chat y perfil, pruebas de carga (30 reqs en 24.5ms) y suite automatizada."),
        ("Luis Alexander Rivera Alvarez", "QA Lead / Database & Security Dev", "14 h",
         "Diseño y ejecución del Plan de QA en Excel (16 TCs), verificación de esquemas PostgreSQL 18 ACID, tipos ENUM, índices B-Tree y pruebas de borrado en cascada (ON DELETE CASCADE)."),
        ("Alberto Jose Velazquez Paz", "Frontend Lead / Desktop UI & QA Tester", "16 h",
         "Implementación de la interfaz de escritorio en Flutter, integración de StateNotifier con Riverpod, diseño de pruebas nativas en macOS (flutter test) y validación cruzada de casos de prueba."),
        ("Luis Angel Zuniga Menjivar", "Backend Dev / API Security & Conformance", "14 h",
         "Auditoría de seguridad en endpoints REST, validación de políticas Fail-Closed en RBAC (HTTP 403), verificación de tokens JWT manipulados (HTTP 401) y testing de concurrencia."),
        ("Ricardo Alberto Mendiola Hernandez", "Dev / Chat Persistente & Perfil Lead", "15 h",
         "Desarrollo del módulo de Chat persistente (IRL-WKS-US-03) con ordenamiento created_at ASC y módulo de personalización de perfil con paleta de 8 colores de avatar y cambio Argon2id (IRL-IAM-US-05)."),
        ("Victor Arnoldo Iglesias Sandoval", "Dev / Reuniones & Servicios Síncronos", "15 h",
         "Desarrollo del módulo de programación de reuniones (IRL-WKS-US-04) con timestamps UTC ISO 8601, selector de duración y enlace directo a Google Meet."),
        ("Jose Luis Fuentes Ochoa", "Dev / Subgrupos & Organización de Nodos", "14 h",
         "Desarrollo del módulo de Subgrupos de nodo (IRL-WKS-US-02), lógica de subgrupos públicos/privados, asignación de creador y ciclo atómico de membresías (Join/Leave).")
    ]
    
    for nom, rol, hrs, det in contributions:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_after = Pt(2)
        r1 = p_c.add_run(f"• {nom} ")
        r1.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = navy
        
        r2 = p_c.add_run(f"({rol} — {hrs})\n")
        r2.bold = True
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = teal
        
        r3 = p_c.add_run(f"  Actividades y Evidencias: {det}")
        r3.font.size = Pt(8.5)
        r3.font.color.rgb = slate_dark
        
    doc.add_paragraph()
    doc.add_heading(level=1).add_run("10. Dictamen Final de Cierre y Aprobación de QA").font.color.rgb = navy
    doc.add_paragraph(
        "El comité de QA y Arquitectura de InnovaSoft certifica que IronLink cumple al 100% con los criterios de aceptación, "
        "DoR, DoD, rendimiento asíncrono y estabilidad en macOS Desktop. El Sprint 2 queda formalmente aprobado y cerrado con éxito."
    )

    output_file = os.path.join(OUTPUT_DIR, "Semana 20 - Ejecución y Cierre del Plan de QA - Sprint 2.docx")
    doc.save(output_file)
    print(f"✅ Informe QA Word exhaustivo generado en: {output_file}")

if __name__ == "__main__":
    build_comprehensive_qa_report()
