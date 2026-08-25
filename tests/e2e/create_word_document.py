import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Sets cell padding in dxa."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout(doc, text, title="NOTA"):
    """Adds a stylish callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F0FDF4") # subtle mint/green background
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="00BFA5"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run_t = p.add_run(f"📌 {title}: ")
    run_t.bold = True
    run_t.font.name = "Arial"
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = RGBColor(0, 150, 136)
    
    run_b = p.add_run(text)
    run_b.font.name = "Arial"
    run_b.font.size = Pt(9.5)
    run_b.font.color.rgb = RGBColor(51, 65, 85)
    doc.add_paragraph()

def build_document():
    doc = docx.Document()
    
    # Page setup (margins: 1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    navy = RGBColor(0, 21, 36)
    teal = RGBColor(0, 191, 165)
    slate_dark = RGBColor(30, 41, 59)
    slate_sub = RGBColor(100, 116, 139)
    
    # ─── PORTADA / CABECERA ────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(4)
    
    run_tag = title_p.add_run("PROYECTO IRONLINK — INGENIERÍA DE SOFTWARE 2\n")
    run_tag.bold = True
    run_tag.font.name = "Arial"
    run_tag.font.size = Pt(11)
    run_tag.font.color.rgb = teal
    
    run_main = title_p.add_run("Especificación Integral de Flujos de Trabajo\ny Reporte de Pruebas Web E2E")
    run_main.bold = True
    run_main.font.name = "Arial"
    run_main.font.size = Pt(22)
    run_main.font.color.rgb = navy
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(20)
    run_sub = sub_p.add_run("Ciclo 2-2026 • Arquitectura Multiplataforma (Flutter Web + Axum Rust + PostgreSQL 18)")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(10.5)
    run_sub.font.color.rgb = slate_sub
    
    # Separator
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(16)
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="00BFA5"/></w:pBdr>')
    sep._p.get_or_add_pPr().append(pBdr)

    # ─── SECCIÓN 1: RESUMEN EJECUTIVO ──────────────────────────────────────
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("1. Resumen Ejecutivo del Testeo Web")
    r1.font.name = "Arial"
    r1.font.size = Pt(16)
    r1.font.color.rgb = navy
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(
        "Se ejecutó una sesión completa de verificación y navegación automatizada de punta a punta (E2E) "
        "sobre el cliente Web de IronLink (http://localhost:3000), validando la integración en tiempo real "
        "con el servidor Backend en Rust Axum (http://127.0.0.1:8080) y la base de datos PostgreSQL 18 (BD: IronLink). "
        "Todos los flujos críticos de la plataforma fueron ejecutados satisfactoriamente sin bloqueos de red, errores de CORS ni fallas de autorización."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = slate_dark
    
    # Tabla de Resultados
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ["ID Flujo", "Funcionalidad Evaluada", "Módulo / Capa", "Resultado"]
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "001524")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.name = "Arial"
        
    results_data = [
        ("FL-01", "Renderizado y validación de formulario de Login", "IAM / Frontend", "✅ APROBADO"),
        ("FL-02", "Registro de cuenta con hashing Argon2id y estado PENDING", "IAM / Backend", "✅ APROBADO"),
        ("FL-03", "Solicitud y envío de Verificación OTP / Enlace Mágico", "IAM / Mailer SMTP", "✅ APROBADO"),
        ("FL-04", "Validación de OTP y activación de usuario a ACTIVE", "IAM / Postgres", "✅ APROBADO"),
        ("FL-05", "Inicio de sesión, emisión JWT y guardado en SecureVault", "Seguridad / Vault", "✅ APROBADO"),
        ("FL-06", "Tablero Principal (Dashboard) con barra y navegación lateral", "Nodos / UI", "✅ APROBADO"),
        ("FL-07", "Creación de nuevo Nodo y generación de token de 32 chars", "Nodos / Backend", "✅ APROBADO"),
        ("FL-08", "Ingreso a la Sala de Chat (# Laboratorio de Ingenieria)", "Chat / GoRouter", "✅ APROBADO"),
        ("FL-09", "Publicación y renderizado persistente de mensajes de chat", "Chat / Postgres", "✅ APROBADO"),
        ("FL-10", "Diálogo de configuración, roles (OWNER/MEMBER) y baneos", "RBAC / Moderación", "✅ APROBADO"),
        ("FL-11", "Diálogo de integración para unirse a nodo por token", "Nodos / Frontend", "✅ APROBADO"),
    ]
    
    for row_idx, data in enumerate(results_data):
        row = table.add_row()
        bg = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.runs[0].font.name = "Arial"
            p.runs[0].font.size = Pt(9)
            if col_idx == 3:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(16, 185, 129)
            else:
                p.runs[0].font.color.rgb = slate_dark
                
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    
    # ─── SECCIÓN 2: ARQUITECTURA GLOBAL ───────────────────────────────────
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("2. Arquitectura Global del Sistema IronLink")
    r2.font.name = "Arial"
    r2.font.size = Pt(16)
    r2.font.color.rgb = navy
    
    p = doc.add_paragraph()
    r = p.add_run(
        "IronLink está diseñado con una arquitectura modular y desacoplada que separa claramente la interfaz de usuario, "
        "la gestión de estado reactiva, los servicios de red con seguridad reforzada y la persistencia relacional asíncrona:"
    )
    r.font.name = "Arial"
    r.font.size = Pt(10)
    
    diag_path = "/Users/ludwin/Developer/ironlink_workspace/tests/e2e/diagrams/diag_01_architecture.png"
    if os.path.exists(diag_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        doc.add_picture(diag_path, width=Inches(6.2))
        
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        rc = cap.add_run("Figura 1: Diagrama de Arquitectura Global — Frontend Flutter, Backend Axum y Base de Datos PostgreSQL.")
        rc.font.italic = True
        rc.font.size = Pt(8.5)
        rc.font.color.rgb = slate_sub

    add_callout(doc, "La comunicación entre Flutter y el backend en Rust se realiza mediante peticiones HTTP REST con cabeceras de autorización Bearer JWT gestionadas por interceptores automáticos de Dio.", "PATRÓN DE RED")

    # ─── SECCIÓN 3: FLUJO 1 - REGISTRO E IDENTIDAD ────────────────────────
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run("3. Flujo 1 — Registro de Usuarios e Identidad (IAM)")
    r3.font.name = "Arial"
    r3.font.size = Pt(16)
    r3.font.color.rgb = navy
    
    p = doc.add_paragraph()
    r = p.add_run(
        "El proceso de registro asegura que ninguna cuenta pueda interactuar con el sistema sin antes someterse a una verificación de identidad. "
        "El backend valida exhaustivamente la unicidad del correo y teléfono, genera un hash Argon2id con salt criptográfico seguro y asigna el estado PENDING."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10)
    
    diag_path = "/Users/ludwin/Developer/ironlink_workspace/tests/e2e/diagrams/diag_02_registration.png"
    if os.path.exists(diag_path):
        doc.add_picture(diag_path, width=Inches(6.2))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        rc = cap.add_run("Figura 2: Diagrama de Pasos del Flujo de Registro de Usuario.")
        rc.font.italic = True
        rc.font.size = Pt(8.5)
        rc.font.color.rgb = slate_sub

    p_scr = doc.add_paragraph()
    p_scr.add_run("Evidencia en la Aplicación Web:").bold = True
    scr_path = "/Users/ludwin/.gemini/antigravity-cli/brain/f4100e48-d450-4d4b-b4a8-e78fae134a31/screenshots/02_register_page.png"
    if os.path.exists(scr_path):
        doc.add_picture(scr_path, width=Inches(5.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(14)
        rc = cap.add_run("Captura 1: Formulario de Registro de IronLink con Stepper (Paso 1: Datos).")
        rc.font.italic = True
        rc.font.size = Pt(8.5)
        rc.font.color.rgb = slate_sub

    # ─── SECCIÓN 4: FLUJO 2 - VERIFICACIÓN DE CUENTA ──────────────────────
    h4 = doc.add_heading(level=1)
    r4 = h4.add_run("4. Flujo 2 — Verificación de Cuenta por Doble Canal")
    r4.font.name = "Arial"
    r4.font.size = Pt(16)
    r4.font.color.rgb = navy
    
    p = doc.add_paragraph()
    r = p.add_run(
        "IronLink implementa un esquema flexible de verificación en dos modalidades:\n"
        "1. Código OTP de 6 dígitos: El usuario recibe un código numérico temporal enviado vía SMTP por Lettre y lo ingresa en pantalla.\n"
        "2. Enlace Mágico (Magic Link): El usuario recibe un enlace directo con token SHA-256 de 64 caracteres. Al hacer clic, el backend valida el token de forma desasistida.\n\n"
        "Una vez verificado el token, el estado en la base de datos cambia automáticamente a ACTIVE."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10)
    
    diag_path = "/Users/ludwin/Developer/ironlink_workspace/tests/e2e/diagrams/diag_03_verification.png"
    if os.path.exists(diag_path):
        doc.add_picture(diag_path, width=Inches(6.2))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        rc = cap.add_run("Figura 3: Diagrama de Verificación de Cuenta (Canales OTP y Enlace por Correo).")
        rc.font.italic = True
        rc.font.size = Pt(8.5)
        rc.font.color.rgb = slate_sub

    scr_v1 = "/Users/ludwin/.gemini/antigravity-cli/brain/f4100e48-d450-4d4b-b4a8-e78fae134a31/screenshots/03_verification_page.png"
    scr_v2 = "/Users/ludwin/.gemini/antigravity-cli/brain/f4100e48-d450-4d4b-b4a8-e78fae134a31/screenshots/04_verification_success_page.png"
    if os.path.exists(scr_v1) and os.path.exists(scr_v2):
        doc.add_picture(scr_v1, width=Inches(5.0))
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        doc.add_picture(scr_v2, width=Inches(5.0))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(14)
        rc = cap.add_run("Capturas 2 y 3: Pantalla de Selección de Método de Verificación y Confirmación de Cuenta Activada.")
        rc.font.italic = True
        rc.font.size = Pt(8.5)
        rc.font.color.rgb = slate_sub

    # ─── SECCIÓN 5: FLUJO 3 - AUTENTICACIÓN Y SESIONES JWT ─────────────────
    h5 = doc.add_heading(level=1)
    r5 = h5.add_run("5. Flujo 3 — Autenticación, JWT y Manejo de Sesión")
    r5.font.name = "Arial"
    r5.font.size = Pt(16)
    r5.font.color.rgb = navy
    
    p = doc.add_paragraph()
    r = p.add_run(
        "La autenticación en IronLink sigue las mejores prácticas de seguridad para aplicaciones empresariales:\n"
        "• Access Token (JWT): De vida corta (15 a 60 minutos), firmado con HMAC-SHA256, contiene claims de usuario y rol.\n"
        "• Refresh Token (UUID): Persistido en PostgreSQL con expiración y revocabilidad ante cierre de sesión.\n"
        "• SecureVault: Encapsula el almacenamiento seguro utilizando DPAPI en Windows, Keychain en macOS y LocalStorage seguro en Web."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10)
    
    diag_path = "/Users/ludwin/Developer/ironlink_workspace/tests/e2e/diagrams/diag_04_auth_jwt.png"
    if os.path.exists(diag_path):
        doc.add_picture(diag_path, width=Inches(6.2))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        rc = cap.add_run("Figura 4: Diagrama de Flujo de Autenticación, Emisión y Persistencia de JWT.")
        rc.font.italic = True
        rc.font.size = Pt(8.5)
        rc.font.color.rgb = slate_sub

    scr_login = "/Users/ludwin/.gemini/antigravity-cli/brain/f4100e48-d450-4d4b-b4a8-e78fae134a31/screenshots/01_login_page.png"
    if os.path.exists(scr_login):
        doc.add_picture(scr_login, width=Inches(5.2))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(14)
        rc = cap.add_run("Captura 4: Pantalla de Inicio de Sesión de IronLink con Paleta Oficial Tech Navy & Mint.")
        rc.font.italic = True
        rc.font.size = Pt(8.5)
        rc.font.color.rgb = slate_sub

    # ─── SECCIÓN 6: FLUJO 4 - GESTIÓN DE NODOS Y COLABORACIÓN ─────────────
    h6 = doc.add_heading(level=1)
    r6 = h6.add_run("6. Flujo 4 — Gestión de Nodos y Colaboración (Workspaces)")
    r6.font.name = "Arial"
    r6.font.size = Pt(16)
    r6.font.color.rgb = navy
    
    p = doc.add_paragraph()
    r = p.add_run(
        "Un Nodo en IronLink constituye una sala o espacio de trabajo colaborativo. "
        "Al crear un nodo, el servidor genera un token_acceso criptográfico de 32 caracteres hexadecimales y registra al creador "
        "con rol OWNER / ADMIN en la tabla nodo_miembros. "
        "Cualquier usuario con el token puede unirse instantáneamente al nodo."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10)
    
    diag_path = "/Users/ludwin/Developer/ironlink_workspace/tests/e2e/diagrams/diag_05_nodos_workspace.png"
    if os.path.exists(diag_path):
        doc.add_picture(diag_path, width=Inches(6.2))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        rc = cap.add_run("Figura 5: Diagrama de Creación, Invitación y Unión a Nodos de Trabajo.")
        rc.font.italic = True
        rc.font.size = Pt(8.5)
        rc.font.color.rgb = slate_sub

    scr_dash = "/Users/ludwin/.gemini/antigravity-cli/brain/f4100e48-d450-4d4b-b4a8-e78fae134a31/screenshots/07_nodos_list_updated.png"
    scr_det = "/Users/ludwin/.gemini/antigravity-cli/brain/f4100e48-d450-4d4b-b4a8-e78fae134a31/screenshots/10_nodo_details_dialog.png"
    if os.path.exists(scr_dash) and os.path.exists(scr_det):
        doc.add_picture(scr_dash, width=Inches(5.0))
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        doc.add_picture(scr_det, width=Inches(5.0))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(14)
        rc = cap.add_run("Capturas 5 y 6: Tablero con Nodo Activo creado y Diálogo Modal de Configuración y Miembros.")
        rc.font.italic = True
        rc.font.size = Pt(8.5)
        rc.font.color.rgb = slate_sub

    # ─── SECCIÓN 7: FLUJO 5 - CHAT PERSISTENTE ─────────────────────────────
    h7 = doc.add_heading(level=1)
    r7 = h7.add_run("7. Flujo 5 — Chat Persistente en Canales de Nodo")
    r7.font.name = "Arial"
    r7.font.size = Pt(16)
    r7.font.color.rgb = navy
    
    p = doc.add_paragraph()
    r = p.add_run(
        "El módulo de mensajería asocia cada mensaje a un nodo y a un usuario autenticado. "
        "El cliente Flutter carga el historial completo ordenado cronológicamente y permite enviar nuevos mensajes "
        "con refresco reactivo instantáneo mediante Riverpod."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10)
    
    diag_path = "/Users/ludwin/Developer/ironlink_workspace/tests/e2e/diagrams/diag_06_chat_messaging.png"
    if os.path.exists(diag_path):
        doc.add_picture(diag_path, width=Inches(6.2))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        rc = cap.add_run("Figura 6: Diagrama de Flujo de Mensajería Persistente en Canales.")
        rc.font.italic = True
        rc.font.size = Pt(8.5)
        rc.font.color.rgb = slate_sub

    scr_chat = "/Users/ludwin/.gemini/antigravity-cli/brain/f4100e48-d450-4d4b-b4a8-e78fae134a31/screenshots/09_nodo_chat_message_sent.png"
    if os.path.exists(scr_chat):
        doc.add_picture(scr_chat, width=Inches(5.2))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(14)
        rc = cap.add_run("Captura 7: Sala de Chat del Nodo con Mensaje Enviado en Tiempo Real y Lista Lateral de Integrantes.")
        rc.font.italic = True
        rc.font.size = Pt(8.5)
        rc.font.color.rgb = slate_sub

    # ─── SECCIÓN 8: MATRIZ REST & RBAC ────────────────────────────────────
    h8 = doc.add_heading(level=1)
    r8 = h8.add_run("8. Matriz de Endpoints REST y Roles RBAC")
    r8.font.name = "Arial"
    r8.font.size = Pt(16)
    r8.font.color.rgb = navy
    
    table_api = doc.add_table(rows=1, cols=4)
    table_api.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_api.rows[0].cells
    api_headers = ["Método", "Endpoint", "Nivel de Autorización", "Descripción Funcional"]
    for i, title in enumerate(api_headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "001524")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.name = "Arial"
        
    api_endpoints = [
        ("POST", "/register", "Público", "Registra usuario con estado PENDING"),
        ("POST", "/login", "Público", "Autentica credenciales y emite tokens JWT + Refresh"),
        ("POST", "/request-verification", "Público", "Emite código OTP de 6 dígitos o enlace por correo"),
        ("POST", "/verify-email", "Público", "Valida código OTP y activa usuario a ACTIVE"),
        ("GET", "/verify-link/{token}", "Público", "Valida enlace web y activa cuenta"),
        ("GET", "/nodos", "MEMBER / ADMIN", "Lista nodos donde el usuario es miembro o creador"),
        ("POST", "/nodos", "MEMBER / ADMIN", "Crea un nuevo nodo y asigna rol ADMIN al creador"),
        ("POST", "/nodos/join/{token}", "MEMBER / ADMIN", "Une al usuario al nodo mediante token de acceso"),
        ("DELETE", "/nodos/{id}", "Creador / ADMIN", "Elimina nodo y todos sus datos en cascada"),
        ("POST", "/nodos/{id}/leave", "Miembro", "Permite abandonar el nodo de trabajo"),
        ("GET", "/nodos/{id}/mensajes", "Miembros", "Obtiene el historial persistente de mensajes"),
        ("POST", "/nodos/{id}/mensajes", "Miembros", "Publica un nuevo mensaje en el canal"),
        ("GET", "/nodos/{id}/miembros", "Miembros", "Lista integrantes y sus respectivos roles"),
        ("POST", "/nodos/{id}/miembros/{uid}/ban", "ADMIN del Nodo", "Banea a un usuario e impide su reingreso"),
        ("DELETE", "/nodos/{id}/baneos/{uid}", "ADMIN del Nodo", "Revoca el baneo de un usuario"),
    ]
    
    for row_idx, data in enumerate(api_endpoints):
        row = table_api.add_row()
        bg = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            p.runs[0].font.name = "Arial"
            p.runs[0].font.size = Pt(8.5)
            if col_idx == 0:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(0, 191, 165) if text in ["GET", "POST"] else RGBColor(239, 68, 68)
            else:
                p.runs[0].font.color.rgb = slate_dark
                
    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    
    # ─── SECCIÓN 9: CONCLUSIONES ──────────────────────────────────────────
    h9 = doc.add_heading(level=1)
    r9 = h9.add_run("9. Conclusiones y Estado para Sprint 2")
    r9.font.name = "Arial"
    r9.font.size = Pt(16)
    r9.font.color.rgb = navy
    
    p = doc.add_paragraph()
    r = p.add_run(
        "1. La plataforma Web de IronLink se encuentra 100% funcional y verificada en todos sus componentes clave del Sprint 1 e inicio del Sprint 2.\n"
        "2. El esquema de seguridad implementado cumple con los criterios de seguridad rigurosa (Argon2id, fail-closed, sesiones JWT revocables y almacenamiento local protegido).\n"
        "3. La base de datos PostgreSQL mantiene integridad referencial en cascada para todos los módulos dependientes (usuarios, tokens, nodos, miembros y mensajes).\n"
        "4. El entorno está listo para continuar con las historias de usuario de Subgrupos (IRL-WKS-US-02), Reuniones (IRL-WKS-US-04) y Perfil de Usuario (IRL-IAM-US-05)."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10)
    
    out_file = "/Users/ludwin/Developer/ironlink_workspace/Flujos_de_Trabajo_IronLink.docx"
    doc.save(out_file)
    print("Word document generated successfully at:", out_file)

if __name__ == "__main__":
    build_document()
